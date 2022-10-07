# Save txt results in a single .docx file
import os
from glob import glob
from docx import Document

folder = "./Results/"

results = sorted(os.listdir(folder))
results = results[1:]


for result in results:
    doc = Document()
    doc.add_heading(result)
    result_path = folder + result + "/"
    txt_files = [x for x in sorted(glob(result_path + "*.txt")) if "_no" in x]
    for txt_file in txt_files:
        with open(txt_file, "r") as f:
            for line in f:
                doc.add_paragraph(line)

    doc.save("./{}.docx".format(result))
